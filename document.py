from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture(
    'me.png', 
    width=Inches(1.5)
)

name = input('what is your name? ')
Phone_number = input('what is your phone number? ')
email = input('what is your email? ')

document.add_paragraph(
    name + ' | ' + Phone_number + ' | ' + email)

#About me 
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# Work experience
document.add_heading('Work Expreience')
p = document.add_paragraph()

company = input('company_name ')
from_date = input('from date')
to_date = input('to date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' ' + to_date + '\n').italic

experience_details = input(
    'Describe your experince at ' +  company)
p.add_run(experience_details)


#more experiences
while True:
    has_more_experience = input(
        'Do you have more experiences? Yes or No' )
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('company_name ')
        from_date = input('from date')
        to_date = input('to date')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' ' + to_date + '\n').italic

        experience_details = input(
            'Describe your experince at ' +  company + ' ')
        p.add_run(experience_details)
    else: 
        break

# Skills
document.add_heading('Skills')
skill = input('Enter Skills')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break





document.save('cv.docx')